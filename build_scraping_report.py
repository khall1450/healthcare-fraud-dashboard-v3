"""Regenerate SCRAPING_REPORT.md and SCRAPING_REPORT.docx from live code + data.

Introspects update.py (FEEDS + scrape_* docstrings), .github/workflows/*.yml
(cron schedules), and data/actions.json (agency coverage counts). Merges the
auto-generated sections into _scraping_report_template.md (hand-written
context + placeholders) to produce the final reports.

Run manually:
    python build_scraping_report.py

Scheduled via .github/workflows/scraping-report.yml (daily 11:30 UTC).
"""
from __future__ import annotations
import io, json, os, re, sys, textwrap
from collections import Counter
from datetime import datetime

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
UPDATE_PY = os.path.join(SCRIPT_DIR, "update.py")
WORKFLOWS_DIR = os.path.join(SCRIPT_DIR, ".github", "workflows")
ACTIONS_JSON = os.path.join(SCRIPT_DIR, "data", "actions.json")
MEDIA_JSON = os.path.join(SCRIPT_DIR, "data", "media.json")
TEMPLATE = os.path.join(SCRIPT_DIR, "_scraping_report_template.md")
OUT_MD = os.path.join(SCRIPT_DIR, "SCRAPING_REPORT.md")
OUT_DOCX = os.path.join(SCRIPT_DIR, "SCRAPING_REPORT.docx")


# ---------------------------------------------------------------------------
# Source-code introspection
# ---------------------------------------------------------------------------

def parse_feeds():
    """Extract the FEEDS list from update.py. Returns list of dicts."""
    src = io.open(UPDATE_PY, encoding="utf-8").read()
    # Find "FEEDS = [" through matching "]"
    m = re.search(r"^FEEDS\s*=\s*\[", src, re.MULTILINE)
    if not m:
        return []
    start = m.end() - 1  # at the '['
    depth, i = 0, start
    while i < len(src):
        if src[i] == "[":
            depth += 1
        elif src[i] == "]":
            depth -= 1
            if depth == 0:
                break
        i += 1
    block = src[start:i + 1]
    # Python-eval each dict line (they're all `{"name": ..., ...}`)
    feeds = []
    for line in block.split("\n"):
        line = line.strip().rstrip(",")
        if line.startswith("{") and line.endswith("}"):
            try:
                # Replace Python None/True/False literals — they're already
                # valid Python; just eval in a restricted namespace
                feed = eval(line, {"None": None, "True": True, "False": False})
                feeds.append(feed)
            except Exception:
                pass
    return feeds


def parse_scrape_functions():
    """Return dict of {scrape_key: {'docstring': str, 'urls': [str]}}."""
    src = io.open(UPDATE_PY, encoding="utf-8").read()
    out = {}
    # Find each `def scrape_X(session):` + its docstring + body
    pattern = re.compile(
        r'def (scrape_\w+)\s*\(session\):\s*\n\s*"""(.*?)"""\s*\n(.*?)(?=\ndef |\Z)',
        re.DOTALL,
    )
    for m in pattern.finditer(src):
        name, doc, body = m.group(1), m.group(2), m.group(3)
        # Normalize docstring: first-paragraph summary
        first_para = doc.strip().split("\n\n")[0]
        first_para = re.sub(r"\s+", " ", first_para).strip()
        # URLs in the function body
        urls = re.findall(r'"(https?://[^"\s]+)"|\'(https?://[^\'\s]+)\'', body)
        urls = [u[0] or u[1] for u in urls]
        # Dedup, keep first 3
        seen, ordered = set(), []
        for u in urls:
            if u not in seen:
                seen.add(u)
                ordered.append(u)
        out[name] = {
            "docstring": first_para,
            "urls": ordered[:3],
        }
    return out


def parse_workflow_crons():
    """Scan .github/workflows/*.yml for cron + name. Returns list."""
    out = []
    if not os.path.isdir(WORKFLOWS_DIR):
        return out
    for fn in sorted(os.listdir(WORKFLOWS_DIR)):
        if not fn.endswith(".yml"):
            continue
        path = os.path.join(WORKFLOWS_DIR, fn)
        text = io.open(path, encoding="utf-8").read()
        # Find name: field (first `name:` at line start)
        name_match = re.search(r"^name:\s*(.+)$", text, re.MULTILINE)
        name = name_match.group(1).strip() if name_match else fn
        # Find cron expressions
        cron_matches = re.findall(r"cron:\s*['\"]([^'\"]+)['\"]", text)
        for cron in cron_matches:
            out.append({
                "file": fn,
                "name": name,
                "cron": cron,
            })
    return out


def agency_counts():
    """Count items by agency in actions.json + media total."""
    with io.open(ACTIONS_JSON, encoding="utf-8") as f:
        a = json.load(f)
    counts = Counter(i.get("agency", "") for i in a.get("actions", []))
    media_total = 0
    try:
        with io.open(MEDIA_JSON, encoding="utf-8") as f:
            m = json.load(f)
        media_total = len(m.get("stories", []))
    except Exception:
        pass
    return counts, media_total


# ---------------------------------------------------------------------------
# Markdown generators for each auto-generated block
# ---------------------------------------------------------------------------

def fmt_cron(cron):
    """Translate '13 9 * * *' to a human-friendly 'Daily 9:13 UTC' style."""
    parts = cron.split()
    if len(parts) != 5:
        return cron
    minute, hour, dom, mon, dow = parts
    # Daily
    if dom == "*" and mon == "*" and dow == "*":
        return f"Daily {hour}:{minute.zfill(2)} UTC"
    # Weekly (single DOW)
    if dom == "*" and mon == "*" and dow.isdigit():
        dow_names = ["Sun", "Mon", "Tue", "Wed", "Thu", "Fri", "Sat"]
        return f"{dow_names[int(dow)]}s {hour}:{minute.zfill(2)} UTC"
    return cron


def gen_schedule_table(crons):
    lines = ["| Workflow | Schedule | File |", "|---|---|---|"]
    for c in crons:
        lines.append(f"| {c['name']} | {fmt_cron(c['cron'])} | `{c['file']}` |")
    return "\n".join(lines)


def gen_feeds_list(feeds, scrape_fns):
    """Generate a full agency-grouped feed list."""
    # Group by agency
    by_agency = {}
    for f in feeds:
        ag = f.get("agency", "Unknown")
        by_agency.setdefault(ag, []).append(f)

    out = []
    order = ["DOJ", "HHS-OIG", "CMS", "HHS", "Congress", "White House",
             "GAO", "MedPAC", "MACPAC", "Treasury"]
    # Put known ones in preferred order, then any leftovers alphabetized
    seen = set()
    ordered_agencies = []
    for ag in order:
        if ag in by_agency:
            ordered_agencies.append(ag)
            seen.add(ag)
    for ag in sorted(k for k in by_agency if k not in seen):
        ordered_agencies.append(ag)

    for ag in ordered_agencies:
        out.append(f"### {ag}\n")
        for f in by_agency[ag]:
            name = f.get("name", "?")
            enabled = f.get("enabled", True)
            scrape_key = f.get("scrape")
            url = f.get("url")
            src_type = f.get("source_type", "official")
            out.append(f"- **`{name}`** — *{src_type}* {'(enabled)' if enabled else '(**disabled**)'}")
            if scrape_key:
                fn = f"scrape_{scrape_key}"
                # Try direct match first, then some known suffix variants
                if fn not in scrape_fns and fn + "_page" in scrape_fns:
                    fn = fn + "_page"
                if fn in scrape_fns:
                    doc = scrape_fns[fn]["docstring"]
                    out.append(f"  - Function: `{fn}()`")
                    out.append(f"  - {doc}")
                    urls = scrape_fns[fn]["urls"]
                    if urls:
                        for u in urls[:2]:
                            out.append(f"  - URL: {u}")
                else:
                    out.append(f"  - Function: `scrape_{scrape_key}()` _(function not found in update.py)_")
            elif url:
                out.append(f"  - Method: RSS feed")
                out.append(f"  - URL: {url}")
            else:
                out.append(f"  - No scraper wired up")
        out.append("")
    return "\n".join(out)


def gen_coverage_table(counts, media_total):
    lines = ["| Source | Items |", "|---|---|"]
    for ag, n in sorted(counts.items(), key=lambda x: -x[1]):
        lines.append(f"| {ag} | {n} |")
    lines.append(f"| Media (manual) | {media_total} |")
    total = sum(counts.values()) + media_total
    lines.append(f"| **Total** | **{total}** |")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Template merge
# ---------------------------------------------------------------------------

def render_template():
    feeds = parse_feeds()
    scrape_fns = parse_scrape_functions()
    crons = parse_workflow_crons()
    counts, media_total = agency_counts()

    now = datetime.now().strftime("%Y-%m-%d %H:%M UTC")
    replacements = {
        "{{GENERATED_AT}}": now,
        "{{AUTO_FEEDS}}": gen_feeds_list(feeds, scrape_fns),
        "{{AUTO_SCHEDULE}}": gen_schedule_table(crons),
        "{{AUTO_COVERAGE}}": gen_coverage_table(counts, media_total),
        "{{AUTO_FEED_COUNT}}": str(len(feeds)),
        "{{AUTO_SCRAPER_COUNT}}": str(len(scrape_fns)),
    }

    if not os.path.exists(TEMPLATE):
        print(f"ERROR: template not found at {TEMPLATE}")
        sys.exit(1)

    tmpl = io.open(TEMPLATE, encoding="utf-8").read()
    for key, val in replacements.items():
        tmpl = tmpl.replace(key, val)

    with io.open(OUT_MD, "w", encoding="utf-8") as f:
        f.write(tmpl)
    print(f"Wrote {OUT_MD}")


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------

def make_docx():
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        print("python-docx not installed; skipping .docx generation")
        return

    md = io.open(OUT_MD, encoding="utf-8").read()
    lines = md.split("\n")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    def inline(paragraph, text):
        pattern = re.compile(r"(\*\*[^*]+\*\*)|(`[^`]+`)|(\[[^\]]+\]\([^)]+\))")
        pos = 0
        for m in pattern.finditer(text):
            if m.start() > pos:
                paragraph.add_run(text[pos:m.start()])
            tok = m.group()
            if tok.startswith("**"):
                r = paragraph.add_run(tok[2:-2]); r.bold = True
            elif tok.startswith("`"):
                r = paragraph.add_run(tok[1:-1])
                r.font.name = "Consolas"; r.font.size = Pt(10)
            elif tok.startswith("["):
                m2 = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
                if m2:
                    paragraph.add_run(m2.group(1))
            pos = m.end()
        if pos < len(text):
            paragraph.add_run(text[pos:])

    i = 0
    in_code = False
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = "Consolas"; r.font.size = Pt(9)
            i += 1
            continue
        if line.strip() in ("---", "***", "___"):
            doc.add_paragraph("_" * 50)
            i += 1
            continue
        mh = re.match(r"^(#{1,6})\s+(.+)$", line)
        if mh:
            lvl = len(mh.group(1))
            text = mh.group(2).strip()
            doc.add_heading(text, level=0 if lvl == 1 else min(lvl - 1, 4))
            i += 1
            continue
        # Table
        if line.strip().startswith("|") and i + 1 < len(lines) and "|--" in lines[i + 1]:
            rows = []
            j = i
            while j < len(lines) and lines[j].strip().startswith("|"):
                if re.match(r"^\s*\|[\s:-]+\|", lines[j]):
                    j += 1
                    continue
                rows.append([c.strip() for c in lines[j].strip().strip("|").split("|")])
                j += 1
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                t.style = "Light Grid Accent 1"
                for ri, row in enumerate(rows):
                    for ci, ct in enumerate(row):
                        if ci < len(t.rows[ri].cells):
                            cell = t.rows[ri].cells[ci]
                            cell.text = ""
                            p = cell.paragraphs[0]
                            inline(p, ct)
                            if ri == 0:
                                for rn in p.runs:
                                    rn.bold = True
                doc.add_paragraph()
            i = j
            continue
        # Bullets
        if re.match(r"^\s*[-*]\s", line):
            while i < len(lines) and re.match(r"^\s*[-*]\s", lines[i]):
                m2 = re.match(r"^(\s*)[-*]\s+(.+)$", lines[i])
                indent_chars = len(m2.group(1))
                text = m2.group(2)
                try:
                    p = doc.add_paragraph(style="List Bullet" if indent_chars < 2 else "List Bullet 2")
                except KeyError:
                    p = doc.add_paragraph(style="List Bullet")
                p.text = ""
                inline(p, text)
                i += 1
            continue
        # Numbers
        if re.match(r"^\s*\d+\.\s", line):
            while i < len(lines) and re.match(r"^\s*\d+\.\s", lines[i]):
                m2 = re.match(r"^\s*\d+\.\s+(.+)$", lines[i])
                try:
                    p = doc.add_paragraph(style="List Number")
                except KeyError:
                    p = doc.add_paragraph()
                p.text = ""
                inline(p, m2.group(1))
                i += 1
            continue
        # Regular
        if line.strip():
            p = doc.add_paragraph()
            inline(p, line)
        i += 1

    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")


# ---------------------------------------------------------------------------

if __name__ == "__main__":
    render_template()
    make_docx()
